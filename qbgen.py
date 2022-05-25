from openpyxl import load_workbook
from openpyxl import Workbook
import random
import string
import io
from PIL import Image
from docx import Document
from docx.shared import Inches
import os

import shutil
questions_with_images=[]
sheet_images={}

class SET_A:
    def __init__(self,data,limit):
        self.question_array_offset = 0
        self.question_array = data
        self.question_array_length = len(data)
        self.question_limit = limit
        self.question_count = 0
        self.random = random
        self.question_bank = []
        self.module = 0
        self.empty = 0
        
        
        self.questions_with_pic_array=[]
    
        
        self.compute_module()
        self.generate_question()   
    

    def compute_module(self):
        data = []
        for q in self.question_array:
            data.append(q[0])
        data.sort()
        data = set(data)
        self.module = len(data)
    def find_question(self):
        question_selector = self.random.randint(self.question_array_offset,self.question_array_length-1)
        question = self.question_array[question_selector]
        return question
    def select_question(self):
        question = self.find_question()
        if question not in self.question_bank:
            found = 0
            try:
                for q in self.question_bank:
                    if(self.question_limit>self.module):
                        self.empty+=1
                    else:
                        if q[0]== question[0]:
                            found += 1
                    if q[1] == question[1] and q[0]== question[0]:
                        found += 1
                    if q[2] == question[2]:
                        found +=1
            except:
                print(0)
            if(found==0):
                self.question_bank.append(question)
                self.question_count+=1

                    
    def generate_question(self):
        while(self.question_limit>self.question_count):
            self.select_question()
    #output generation
    def questions(self):
        return self.question_bank
    def questions_only(self):
        data = []
        for d in self.question_bank:
            data.append(d[2])
        return data
    
    
    
def create_folder(fp):
    directory = "questionfiles"
    parent_dir =fp
    path = os.path.join(parent_dir, directory)
    try: 
        os.mkdir(path) 
    except OSError as error:
        print(error)  
    
def delete_folder(fp):
    directory = "questionfiles"
    parent_dir =fp
    path = os.path.join(parent_dir, directory)
    print(path)
    try:
        shutil.rmtree(path)
    except OSError:
        print ("Deletion of the directory %s failed" % path)
    else:
        print ("Successfully deleted the directory %s" % path)
    
    
def delete_file(fp):
    directory = "demo.docx"
    parent_dir =fp
    path = os.path.join(parent_dir, directory)
    try:
        os.remove(path)
    except OSError:
        print ("Deletion of the directory %s failed" % path)
    else:
        print ("Successfully deleted the directory %s" % path)



def image_in(section):
    return section in sheet_images

def getfilename(j,fp):
    directory = "questionfiles"
    parent_dir =fp
    path = os.path.join(parent_dir, directory)
    filename=str(j)+".png"
    path = os.path.join(path,filename)
    return path
    
    
def load_question_images(fp):
    # loads the images in the sheet
    workbook=load_workbook(filename= fp)
    for i in workbook.sheetnames:
        sheet=workbook[i]
        global questions_with_images
        global sheet_images
        for image in sheet._images:
            row = image.anchor._from.row + 1
            col = string.ascii_uppercase[image.anchor._from.col]
            sheet_images[f'{sheet}{col}{row}'] = image._data

def initialiase_section(section,fp):
    
            
    workbook=load_workbook(filename= fp)
    ws=workbook[section]
    my_array = []
    sheet=ws
    s=ws.max_row
    for i in range(10,s):
        temp_array=[]
        getmodule =ws.cell(i,1)
        currentunit=ws.cell(i,2)
        currentquestion=ws.cell(i,3)
        try:

            temp_array.insert(0,int(getmodule.value))
            #print(ws,i,getmodule.value,ws.max_row)
        except:
            print("This except block is executed because there is an empty row in between the rows in that case kindly check the excel file or because")
            print("openpyxl made a wrong calculation about the max number of rows in ",ws," so the number of rows of the current worksheet must be entered manually")
            #s=int(input("please enter the max number of rows: "))
            break


        temp_array.insert(1,int(currentunit.value))
        temp_array.insert(2,str(currentquestion.value))
        section_="{sheet}{cell}".format(sheet=ws,cell=currentquestion.coordinate)
        if image_in(section_):
            questions_with_images.append(currentquestion.value)
            

        my_array.append(temp_array)
    return my_array
        
def savequestion(s1,s2,s3):
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    fp=os.getcwd()
    document = Document()
    document.add_heading('Model question paper', 0)
    #f=open("writefile.txt","w")
    #f.write("SECTION-A \n")
    questionnumber=1
    p = document.add_paragraph()
    
    
    r=p.add_run()
    questionnumber=1
    for i in range(len(s1)):
        r.add_break()
        r.add_text("%d. %s [%d,%d]\n"%(questionnumber,s1[i][2],s1[i][0],s1[i][1]))
        #p=document.add_paragraph()
        if s1[i][2] in questions_with_images:
            #print("question")
            for j in range(len(questions_with_images)):
                if questions_with_images[j] == s1[i][2]:
                    #print("question")
                    cell=sheet_images[list(sheet_images)[j]]
                    image=io.BytesIO(cell())
                    question_=Image.open(image)
                    # save file in the directory and use it as argument for the r.add_picture()
                    question_.save(getfilename(j,fp))
                    r.add_picture(getfilename(j,fp))
        r.add_break()
        #p=document.add_paragraph()
        #f.write("%d. %s [%d,%d]\n"%(questionnumber,s3[i][2],s3[i][0],s3[i][1]))
        questionnumber+=1
        
    for i in range(len(s2)):
        r.add_break()
        r.add_text("%d. %s [%d,%d]\n"%(questionnumber,s2[i][2],s2[i][0],s2[i][1]))
        if s2[i][2] in questions_with_images:
            #print("question")
            for j in range(len(questions_with_images)):
                if questions_with_images[j] == s2[i][2]:
                    #print("question")
                    cell=sheet_images[list(sheet_images)[j]]
                    image=io.BytesIO(cell())
                    question_=Image.open(image)
                    # save file in the directory and use it as argument for the r.add_picture()
                    question_.save(getfilename(j,fp))
                    r.add_picture(getfilename(j,fp))
        r.add_break()
        #p=document.add_paragraph()
        #f.write("%d. %s [%d,%d]\n"%(questionnumber,s3[i][2],s3[i][0],s3[i][1]))
        questionnumber+=1
        
    
    
    for i in range(len(s3)):
        r.add_break()
        r.add_text("%d. %s [%d,%d]\n"%(questionnumber,s3[i][2],s3[i][0],s3[i][1]))
        #p=document.add_paragraph()
        if s3[i][2] in questions_with_images:
            #print("question")
            for j in range(len(questions_with_images)):
                if questions_with_images[j] == s3[i][2]:
                    #print("question")
                    cell=sheet_images[list(sheet_images)[j]]
                    image=io.BytesIO(cell())
                    question_=Image.open(image)
                    # save file in the directory and use it as argument for the r.add_picture()
                    question_.save(getfilename(j,fp))
                    r.add_picture(getfilename(j,fp))
        r.add_break()
        #p=document.add_paragraph()
        #f.write("%d. %s [%d,%d]\n"%(questionnumber,s3[i][2],s3[i][0],s3[i][1]))
        questionnumber+=1
    document.save("E:\\Ravana\\workstation\\general\\Coronis\\GameOFThreads\\static\\"+'demo.docx')
    #f.close()
    

def deleteStaticFiles():
    try:
        for file in os.listdir("E:\\Ravana\\workstation\\general\\Coronis\\GameOFThreads\\static\\"):
          os.remove("E:\\Ravana\\workstation\\general\\Coronis\\GameOFThreads\\static\\"+file)

    except:
        pass
    

def acceptPath(filepath):
    fp = filepath
    #fp=input("Enter the full filepath to the questionbank file - use / when entering the filepath: ")
    # cwd = os.getcwd()
    # create_folder(cwd)
    # parent_dir =cwd
    # file= os.path.join(parent_dir,"demo.docx")
    # if os.path.exists(file):
    #     delete_file(cwd)

    load_question_images(fp)
    trial1=SET_A(data=initialiase_section('section-a',fp),limit=5)
    trial2=SET_A(data=initialiase_section('section-b',fp),limit=7)
    trial3=SET_A(data=initialiase_section('section-c',fp),limit=5)

    savequestion(trial1.questions(),trial2.questions(),trial3.questions())
    #delete_folder(cwd)
    #print("/////////Running the script on the same directory will delete the previous word file//////////")

#file = ""C:\\Users\\jetle\\Desktop\\New_DBMS_QBankDBMS.xlsx""
#acceptPath(file)